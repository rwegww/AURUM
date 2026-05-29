from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs" / "aurum_luong_nghiep_vu_he_thong.docx"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "B7C3D0"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[i] / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[i]))
            tc_w.set(qn("w:type"), "dxa")


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        s = styles[name]
        s.font.name = "Calibri"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("Mô tả luồng nghiệp vụ hệ thống AURUM")
    r.font.name = "Calibri"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run(
        "Tài liệu phân tích dựa trên worktree hiện tại: React/Vite frontend, Express API, Supabase, Netlify serverless và các migration SQL."
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(85, 85, 85)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        set_cell_shading(hdr[i], LIGHT_BLUE)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
    set_table_geometry(table, widths)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_callout(doc: Document, label: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    r = p.add_run(label + ": ")
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    p.add_run(text)
    set_table_geometry(table, [9360])
    doc.add_paragraph()


def build() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_title(doc)

    doc.add_heading("1. Tóm tắt hệ thống", level=1)
    doc.add_paragraph(
        "AURUM là nền tảng học Hóa học dạng game hóa. Hệ thống phục vụ ba nhóm người dùng chính: học sinh, giáo viên và quản trị viên. "
        "Frontend React/Vite điều hướng trải nghiệm học tập, phòng thí nghiệm, đấu trường, thư viện và quản trị. Backend Express gom các module API, xác thực người dùng, áp quyền theo vai trò và ghi dữ liệu vào Supabase."
    )
    add_callout(
        doc,
        "Kết luận nghiệp vụ",
        "Luồng nghiệp vụ trung tâm là: người dùng đăng ký/đăng nhập, hệ thống xác định vai trò, học sinh học theo hành trình và nhận XP/streak/nhiệm vụ, giáo viên quản lý lớp và bài tập, quản trị viên kiểm soát nội dung/người dùng/phản hồi."
    )

    add_table(
        doc,
        ["Vai trò", "Mục tiêu nghiệp vụ", "Khu vực chính trong hệ thống"],
        [
            ["Khách", "Xem thông tin công khai, đăng ký, đăng nhập hoặc gửi yêu cầu tài khoản giáo viên.", "Home, Login, Register, public praises, public leaderboard."],
            ["Học sinh", "Học bài, luyện tập, tham gia lớp, nộp bài, dùng phòng thí nghiệm, đấu trường và theo dõi tiến độ.", "Classroom, Journey, Lab, Arena, Library, My Class, Profile, Missions."],
            ["Giáo viên", "Tạo lớp, giao bài, đăng thông báo/video, chấm điểm, xem thống kê và thông báo lớp.", "Teacher dashboard, Class Manager, Class Detail, Assignment Manager."],
            ["Quản trị viên", "Quản lý người dùng, khóa tài khoản, quản lý bài học, duyệt phản hồi và yêu cầu giáo viên.", "Admin dashboard, users, lessons, journey, feedback."],
            ["Tác vụ hệ thống", "Duy trì nhắc học, streak, heartbeat hoạt động và cron email.", "Heartbeat mỗi phút, cron reminder, Netlify scheduled function."],
        ],
        [1500, 3800, 4060],
    )

    doc.add_heading("2. Biên hệ thống và kiến trúc nghiệp vụ", level=1)
    doc.add_paragraph(
        "Luồng tổng quát đi từ trình duyệt đến React Router, qua AuthProvider để xác định phiên đăng nhập, sau đó gọi Express API. API dùng middleware xác thực, service/model layer và Supabase để đọc ghi dữ liệu."
    )
    add_numbered(
        doc,
        [
            "Người dùng thao tác trên frontend: trang công khai, trang học sinh, trang giáo viên hoặc trang admin.",
            "AuthProvider đọc token trong localStorage hoặc Supabase session, gọi /api/user/profile để lấy hồ sơ và claim session.",
            "Express nhận request tại /api/*, route đúng module nghiệp vụ, kiểm tra token bằng JWT nội bộ hoặc Supabase OAuth.",
            "Model/service layer đọc ghi Supabase; một số luồng dùng Supabase OAuth, Cloudinary upload, SMTP email, Netlify scheduled function.",
            "Frontend cập nhật trạng thái người dùng, XP, level, streak, tiến độ, bài nộp và thông báo theo response API.",
        ],
    )

    doc.add_heading("3. Luồng xác thực và quản lý phiên", level=1)
    add_bullets(
        doc,
        [
            "Đăng ký học sinh tạo bản ghi users, hash mật khẩu, khởi tạo user_progress cho bài học/hóa chất/balancing và trả JWT nội bộ.",
            "Đăng nhập thường kiểm tra username/password, tạo JWT có sessionId, lưu token/sessionId/authType ở localStorage.",
            "Đăng nhập Google đi qua Supabase OAuth; nếu chưa có user tương ứng, middleware tạo user mới với role student.",
            "Yêu cầu giáo viên không tạo tài khoản ngay; hệ thống ghi một feedback type teacher_registration để admin duyệt hoặc từ chối.",
            "Khi lấy profile với claim=true, backend ghi current_session_id; request sau bị chặn nếu sessionId không khớp, hỗ trợ chống đăng nhập song song.",
            "Tài khoản bị khóa bằng admin route sẽ bị middleware auth từ chối trước khi vào nghiệp vụ.",
        ],
    )

    doc.add_heading("4. Luồng học tập của học sinh", level=1)
    add_numbered(
        doc,
        [
            "Học sinh chọn khối/lộ trình ở Classroom hoặc Grade Journey; frontend tải danh sách lessons theo class_id/program_id.",
            "Học sinh đi qua các màn hình intro, story, challenge, quiz và reward của từng lesson.",
            "Khi hoàn thành từng cấp độ, frontend gọi /api/user/lesson-segment với lessonId, level và số sao.",
            "Backend kiểm tra lesson tồn tại, cập nhật lessonStars trong user_progress item_type=balancing, cộng XP lần đầu cho mỗi level.",
            "Khi hoàn thành level3, hệ thống mở khóa lesson, đánh dấu hoàn thành mục tiêu ngày, cập nhật streak và studyPlan.completed.",
            "Nếu học sinh vượt placement test, /api/user/placement-pass ghi passedGrades và cộng 500 XP nếu là lần đầu.",
        ],
    )
    add_table(
        doc,
        ["Dữ liệu", "Bảng/field chính", "Ý nghĩa nghiệp vụ"],
        [
            ["Hồ sơ học sinh", "users", "Role, XP, level, avatar, streak, study_plan, trạng thái online."],
            ["Tiến độ học/lab", "user_progress", "Nguồn chính cho unlocked lessons, unlocked chemicals, balancing progress và lesson stars."],
            ["Nội dung bài học", "lessons", "Module lý thuyết, video, quiz, story, challenge, game, intro video."],
            ["Nhắc học", "users.study_plan", "Mục tiêu học mỗi ngày, emailEnabled, completed, trạng thái nhắc học."],
        ],
        [1800, 2600, 4960],
    )

    doc.add_heading("5. Luồng nhiệm vụ, XP, level và streak", level=1)
    add_bullets(
        doc,
        [
            "Missions được đọc từ bảng missions; tiến độ cá nhân nằm ở user_missions.",
            "Daily missions tự reset theo ngày khi người dùng đọc hoặc cập nhật nhiệm vụ.",
            "claimReward gọi RPC claim_mission_reward để đánh dấu đã nhận thưởng và cộng XP nguyên tử.",
            "Heartbeat mỗi phút tăng active_minutes, today_online_minutes và có thể tự thắp streak khi online đủ 10 phút.",
            "Hoàn thành bài học cũng có thể cập nhật streak; streak dài tạo bonus XP theo công thức trong user route.",
            "Cron reminder quét học sinh có emailEnabled, chưa hoàn thành mục tiêu ngày và gửi email nhắc học/streak qua SMTP.",
        ],
    )

    doc.add_heading("6. Luồng lớp học và bài tập", level=1)
    add_numbered(
        doc,
        [
            "Giáo viên hoặc admin tạo lớp; hệ thống sinh mã lớp duy nhất trong bảng classes.",
            "Học sinh nhập mã lớp; nếu hợp lệ, hệ thống ghi class_members và chặn join trùng.",
            "Giáo viên tạo posts trong lớp: announcement, assignment hoặc video; có thể gán cho toàn lớp hoặc học sinh cụ thể.",
            "Học sinh xem posts theo quyền: bài chung, bài được giao riêng, hoặc bài do chính học sinh gửi.",
            "Bài tập có thể chứa questions parse từ file đề; học sinh nộp answers vào class_assignment_submissions.",
            "Giáo viên xem tiến độ nộp bài, chấm score/feedback và theo dõi notifications từ học sinh mới, tin nhắn, bài nộp và bài sắp hết hạn.",
        ],
    )
    add_table(
        doc,
        ["Bảng", "Đối tượng nghiệp vụ", "Ghi chú quyền truy cập"],
        [
            ["classes", "Lớp học", "Teacher sở hữu lớp; admin xem toàn hệ thống."],
            ["class_members", "Quan hệ học sinh - lớp", "Student chỉ thấy lớp mình tham gia."],
            ["class_posts", "Thông báo, bài tập, video, tin nhắn", "Teacher quản lý lớp; student thấy bài phù hợp target."],
            ["class_assignment_submissions", "Bài nộp", "Student nộp bài của mình; teacher chấm bài lớp mình."],
            ["class_schedules", "Lịch học/meet", "Teacher tạo; thành viên lớp xem."],
        ],
        [2300, 3000, 4060],
    )

    doc.add_heading("7. Luồng phòng thí nghiệm, cân bằng và mô phỏng", level=1)
    add_bullets(
        doc,
        [
            "Lab đọc danh sách hóa chất từ lab_chemicals và phản ứng từ lab_reactions.",
            "Balancing search và lookup theo nodeId đọc balancing_questions.",
            "Tiến độ cân bằng được lưu vào user_progress item_type=balancing, item_id=current.",
            "Mở khóa hóa chất ghi user_progress item_type=chemical; các màn hình lab dùng dữ liệu này để cá nhân hóa trải nghiệm.",
            "Các mô phỏng frontend như density, gas law, molarity, pH scale và formula sim chủ yếu chạy client-side.",
        ],
    )

    doc.add_heading("8. Luồng đấu trường Arena", level=1)
    add_numbered(
        doc,
        [
            "Người chơi tạo phòng hoặc join phòng đang waiting; join dùng RPC join_arena_room để tăng current_players an toàn.",
            "Find-match tìm hoặc tạo phòng theo difficulty/mode.",
            "Khi chơi, frontend lấy câu hỏi từ arena_questions theo difficulty.",
            "Kết quả trận đấu ghi arena_match_history, cập nhật arena_stats trong users và dùng cho leaderboard/my battles.",
            "Người chơi có thể lưu arena avatar riêng qua /api/arena/save-avatar.",
        ],
    )

    doc.add_heading("9. Luồng thư viện tài liệu và phản hồi", level=1)
    add_bullets(
        doc,
        [
            "Library đọc danh sách materials có filter; Material Detail tăng view_count qua RPC increment_material_view.",
            "Người dùng gửi đánh giá tài liệu vào material_feedback; giáo viên/admin có thể trả lời feedback.",
            "Feedback hệ thống gồm góp ý, lỗi, lời khen và yêu cầu đăng ký giáo viên.",
            "Admin có thể xem, resolve, approve praise, duyệt hoặc từ chối teacher registration; public home chỉ đọc feedback praise đã approve.",
        ],
    )

    doc.add_heading("10. Luồng quản trị", level=1)
    add_bullets(
        doc,
        [
            "Admin dashboard tổng hợp số học sinh, số bài học, feedback chưa đọc, phân phối feedback và thống kê XP/level/streak.",
            "User Manager xem danh sách học sinh, khóa/mở tài khoản và xem chi tiết người dùng.",
            "Lesson/Journey Manager tạo, sửa, xóa bài học và cấu trúc lộ trình.",
            "Feedback Manager xử lý phản hồi người dùng, lời khen công khai và yêu cầu tài khoản giáo viên.",
        ],
    )

    doc.add_heading("11. Ma trận module - dữ liệu - API", level=1)
    add_table(
        doc,
        ["Module", "API mount", "Bảng/RPC chính"],
        [
            ["Auth", "/api/auth", "users, feedback, JWT, Supabase OAuth."],
            ["User/Profile", "/api/user", "users, user_progress, feedback, cron mailer."],
            ["Lessons", "/api/lessons, /api/admin/lessons", "lessons, grade_levels."],
            ["Missions", "/api/missions", "missions, user_missions, claim_mission_reward."],
            ["Classes", "/api/classes", "classes, class_members, class_posts, class_assignment_submissions, class_schedules."],
            ["Lab", "/api/lab", "lab_chemicals, lab_reactions, balancing_questions, user_progress."],
            ["Arena", "/api/arena", "arena_rooms, arena_questions, arena_match_history, users.arena_stats, join_arena_room."],
            ["Materials", "/api/materials", "materials, material_feedback, increment_material_view."],
            ["Discussions", "/api/discussions", "lesson_discussions, user_notes, increment_likes."],
            ["Admin", "/api/admin", "users, lessons, feedback, aggregate stats."],
        ],
        [1900, 2500, 4960],
    )

    doc.add_heading("12. Quy tắc nghiệp vụ và điểm kiểm soát", level=1)
    add_bullets(
        doc,
        [
            "Route protected dùng middleware auth; teacher/admin routes kiểm tra role hoặc quyền sở hữu lớp.",
            "Student không được join lớp nếu role không phải student; không được nộp bài không thuộc lớp hoặc không được giao.",
            "Teacher không được quản lý lớp của teacher khác, trừ khi là admin.",
            "Progress hiện ưu tiên user_progress; code vẫn có fallback legacy cho user_unlocked_* nhưng live schema sạch không expose các bảng đó.",
            "Backend service-role đảm nhiệm nhiều thao tác ghi; RLS vẫn giới hạn direct client access ở các bảng public.",
            "Session claim + current_session_id làm điểm kiểm soát chống dùng cùng tài khoản ở nhiều nơi.",
        ],
    )

    doc.add_heading("13. Ghi chú vận hành", level=1)
    add_bullets(
        doc,
        [
            "Frontend build bằng Vite; API chạy Express local hoặc qua Netlify function wrapper.",
            "Analyze route được lazy-load để tránh cold-start crash do thư viện xử lý PDF/DOCX nặng.",
            "Cron reminder chạy local bằng setInterval khi dev và có Netlify scheduled function cho môi trường deploy.",
            "Database hiện đã dọn về schema runtime gồm 23 bảng public và 4 RPC exposed theo kiểm tra live Supabase trước đó.",
        ],
    )

    doc.add_paragraph()
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("AURUM business flow analysis - generated from current codebase").font.size = Pt(8)

    doc.save(OUT)


if __name__ == "__main__":
    build()
    print(OUT)
