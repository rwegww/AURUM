from __future__ import annotations

import json
import os
import re
import urllib.error
import urllib.request
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs" / "aurum_bang_mo_ta_co_so_du_lieu.docx"
ENV_PATH = ROOT / ".env.local"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
LIGHT_GREEN = "E8F4EA"


TABLE_DESCRIPTIONS = {
    "users": ("Nguoi dung", "Ho so tai khoan, vai tro, XP, level, streak, session, study plan va trang thai khoa tai khoan."),
    "grade_levels": ("Khoi/lop", "Danh muc khoi hoc va metadata phuc vu lo trinh bai hoc."),
    "lessons": ("Bai hoc", "Noi dung bai hoc, quiz, story, challenge, game data, video va thong tin lo trinh."),
    "user_progress": ("Tien do hoc tap", "Nguon tien do chinh cho lesson, balancing, chemical unlock va lesson stars."),
    "missions": ("Nhiem vu", "Cau hinh nhiem vu ngay/thang/thanh tich va thuong XP."),
    "user_missions": ("Tien do nhiem vu", "Tien do ca nhan cua tung nguoi dung voi tung mission."),
    "feedback": ("Phan hoi", "Gop y, loi, loi khen cong khai va yeu cau dang ky giao vien."),
    "lab_chemicals": ("Hoa chat lab", "Danh muc chat, cong thuc, mau sac, do kho va metadata mo phong."),
    "lab_reactions": ("Phan ung lab", "Cong thuc phan ung, chat tham gia, san pham va dieu kien."),
    "balancing_questions": ("Cau hoi can bang", "Cau hoi can bang phuong trinh va du lieu node trong skill tree."),
    "arena_questions": ("Cau hoi arena", "Ngan hang cau hoi dau truong theo do kho."),
    "arena_rooms": ("Phong arena", "Phong cho, trang thai, nguoi tao, current players va cau hinh match."),
    "arena_match_history": ("Lich su arena", "Ket qua tran dau va thong ke nguoi choi."),
    "classes": ("Lop hoc", "Lop do giao vien/admin tao, ma lop, mo ta va trang thai."),
    "class_members": ("Thanh vien lop", "Quan he hoc sinh tham gia lop."),
    "class_posts": ("Bai dang lop", "Thong bao, bai tap, video, tin nhan va target hoc sinh."),
    "class_assignment_submissions": ("Bai nop", "Cau tra loi, diem, feedback va trang thai nop bai."),
    "class_schedules": ("Lich lop", "Lich hoc, meeting link va thoi gian cua lop."),
    "materials": ("Tai lieu", "Thu vien tai lieu, noi dung trich xuat, file url, tag, view count."),
    "material_feedback": ("Phan hoi tai lieu", "Danh gia tai lieu va phan hoi cua giao vien/admin."),
    "lesson_discussions": ("Thao luan bai hoc", "Binh luan theo bai hoc, reply thread va luot thich."),
    "user_notes": ("Ghi chu ca nhan", "Ghi chu rieng cua nguoi dung theo bai hoc."),
    "user_activities": ("Nhat ky hoat dong", "Log hanh vi hoc tap de tinh nhac hoc, thong ke va audit nhe."),
}

MODULE_MAP = {
    "Auth/Profile": ["users", "feedback"],
    "Learning": ["grade_levels", "lessons", "user_progress", "lesson_discussions", "user_notes"],
    "Missions": ["missions", "user_missions", "users"],
    "Classes": ["classes", "class_members", "class_posts", "class_assignment_submissions", "class_schedules", "lessons"],
    "Lab": ["lab_chemicals", "lab_reactions", "balancing_questions", "user_progress"],
    "Arena": ["arena_questions", "arena_rooms", "arena_match_history", "users"],
    "Materials": ["materials", "material_feedback", "users"],
    "Activity/Reminder": ["user_activities", "users"],
}

RPCS = {
    "increment_likes": "Tang likes cho lesson_discussions.",
    "increment_material_view": "Tang view_count cho materials khi xem chi tiet tai lieu.",
    "claim_mission_reward": "Danh dau mission da nhan thuong va cong XP cho user.",
    "join_arena_room": "Tang current_players phong arena an toan khi nguoi choi tham gia.",
}


def load_env() -> None:
    if not ENV_PATH.exists():
        return
    for raw in ENV_PATH.read_text(encoding="utf-8", errors="ignore").splitlines():
        line = raw.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, value = line.split("=", 1)
        os.environ.setdefault(key.strip(), value.strip().strip('"').strip("'"))


def fetch_json(url: str, headers: dict[str, str]) -> dict:
    req = urllib.request.Request(url, headers=headers)
    with urllib.request.urlopen(req, timeout=30) as res:
        return json.loads(res.read().decode("utf-8"))


def fetch_live_schema() -> tuple[dict, dict[str, int | None], str]:
    load_env()
    supabase_url = os.environ.get("SUPABASE_URL", "").rstrip("/")
    key = os.environ.get("SUPABASE_SERVICE_ROLE_KEY", "")
    if not supabase_url or not key:
        return {}, {}, "Khong tim thay SUPABASE_URL/SUPABASE_SERVICE_ROLE_KEY trong .env.local."

    headers = {
        "apikey": key,
        "Authorization": f"Bearer {key}",
        "Accept": "application/openapi+json",
    }
    openapi = fetch_json(f"{supabase_url}/rest/v1/", headers)
    definitions = openapi.get("definitions", {})

    row_counts: dict[str, int | None] = {}
    for table in definitions:
        req = urllib.request.Request(
            f"{supabase_url}/rest/v1/{table}?select=*&limit=0",
            headers={
                "apikey": key,
                "Authorization": f"Bearer {key}",
                "Prefer": "count=exact",
                "Range": "0-0",
            },
        )
        try:
            with urllib.request.urlopen(req, timeout=20) as res:
                content_range = res.headers.get("content-range") or res.headers.get("Content-Range") or ""
        except urllib.error.HTTPError as err:
            content_range = err.headers.get("content-range") or err.headers.get("Content-Range") or ""
        match = re.search(r"/(\d+)$", content_range)
        row_counts[table] = int(match.group(1)) if match else None

    return definitions, row_counts, "Live Supabase PostgREST OpenAPI + exact count qua service role."


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[i] / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[i]))
            tc_w.set(qn("w:type"), "dxa")


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color, before, after in [
        ("Heading 1", 15, BLUE, 14, 7),
        ("Heading 2", 12, DARK_BLUE, 10, 5),
        ("Heading 3", 11, DARK_BLUE, 8, 3),
    ]:
        s = styles[name]
        s.font.name = "Calibri"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True


def add_title(doc: Document, source: str, table_count: int) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("Bảng mô tả cơ sở dữ liệu AURUM")
    r.font.name = "Calibri"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run(f"Nguồn phân tích: {source} Tổng số bảng public đang expose: {table_count}.")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(85, 85, 85)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int], header_fill: str = LIGHT_BLUE) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        set_cell_shading(hdr[i], header_fill)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = str(text)
    set_table_geometry(table, widths)
    doc.add_paragraph()


def parse_column(table: str, name: str, meta: dict, required: set[str]) -> list[str]:
    desc = meta.get("description", "") or ""
    flags = []
    if "<pk/>" in desc:
        flags.append("PK")
    fk_match = re.search(r"<fk table='([^']+)' column='([^']+)'/>", desc)
    if fk_match:
        flags.append(f"FK -> {fk_match.group(1)}.{fk_match.group(2)}")
    if name in required:
        flags.append("NOT NULL")
    default = meta.get("default")
    if default is not None:
        flags.append(f"default {default}")
    data_type = meta.get("format") or meta.get("type") or "json"
    return [name, data_type, "; ".join(flags) or "-", column_note(table, name)]


def column_note(table: str, column: str) -> str:
    notes = {
        ("users", "current_session_id"): "Chống đăng nhập song song theo phiên hiện hành.",
        ("users", "study_plan"): "Cấu hình mục tiêu học/ngày và nhắc học.",
        ("users", "arena_stats"): "Thống kê thắng/thua/XP arena.",
        ("user_progress", "item_type"): "Phân loại lesson, chemical, balancing hoặc progress item khác.",
        ("user_progress", "metadata"): "Chứa stars, progress, unlock flags hoặc dữ liệu phụ theo item_type.",
        ("class_posts", "target_student_ids"): "Giới hạn bài đăng/bài tập cho nhóm học sinh cụ thể.",
        ("class_posts", "questions"): "Câu hỏi trích từ file bài tập hoặc nhập thủ công.",
        ("materials", "view_count"): "Được tăng qua RPC increment_material_view.",
        ("feedback", "type"): "Phân loại bug, suggestion, praise, teacher_registration.",
    }
    return notes.get((table, column), "-")


def fk_edges(definitions: dict) -> list[tuple[str, str, str]]:
    edges = []
    for table, spec in definitions.items():
        for column, meta in spec.get("properties", {}).items():
            desc = meta.get("description", "") or ""
            match = re.search(r"<fk table='([^']+)' column='([^']+)'/>", desc)
            if match:
                edges.append((f"{table}.{column}", match.group(1), match.group(2)))
    return sorted(edges)


def build() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    definitions, row_counts, source = fetch_live_schema()
    if not definitions:
        raise RuntimeError(source)

    doc = Document()
    style_document(doc)
    table_names = sorted(definitions)
    add_title(doc, source, len(table_names))

    doc.add_heading("1. Tóm tắt mô hình dữ liệu", level=1)
    doc.add_paragraph(
        "Database public hiện tập trung vào các nhóm nghiệp vụ: tài khoản và phiên, học tập, nhiệm vụ, lớp học, phòng thí nghiệm, arena, thư viện tài liệu, thảo luận và nhật ký hoạt động. "
        "Các bảng legacy/AI/content thử nghiệm không còn xuất hiện trong schema live được PostgREST expose."
    )
    add_table(
        doc,
        ["Nhóm nghiệp vụ", "Bảng sử dụng"],
        [[module, ", ".join(tables)] for module, tables in MODULE_MAP.items()],
        [2600, 6760],
        LIGHT_GREEN,
    )

    doc.add_heading("2. Danh sách bảng", level=1)
    summary_rows = []
    for table in table_names:
        group, description = TABLE_DESCRIPTIONS.get(table, ("Khác", "Chưa có mô tả nghiệp vụ chi tiết."))
        column_count = len(definitions[table].get("properties", {}))
        count = row_counts.get(table)
        summary_rows.append([table, group, column_count, count if count is not None else "n/a", description])
    add_table(doc, ["Bảng", "Nhóm", "Cột", "Rows live", "Mô tả"], summary_rows, [1750, 1600, 700, 900, 4410])

    doc.add_heading("3. Quan hệ khóa ngoại chính", level=1)
    edge_rows = [[src, f"{target_table}.{target_column}", "Ràng buộc tham chiếu"] for src, target_table, target_column in fk_edges(definitions)]
    add_table(doc, ["Cột nguồn", "Tham chiếu", "Ý nghĩa"], edge_rows, [3300, 3000, 3060])

    doc.add_heading("4. RPC/function đang dùng", level=1)
    add_table(doc, ["RPC", "Mục đích nghiệp vụ"], [[name, desc] for name, desc in RPCS.items()], [3000, 6360], LIGHT_GREEN)

    doc.add_heading("5. Mô tả chi tiết từng bảng", level=1)
    for table in table_names:
        group, description = TABLE_DESCRIPTIONS.get(table, ("Khác", "Chưa có mô tả nghiệp vụ chi tiết."))
        row_count = row_counts.get(table)
        doc.add_heading(table, level=2)
        doc.add_paragraph(f"Nhóm: {group}. Rows live: {row_count if row_count is not None else 'n/a'}. {description}")
        spec = definitions[table]
        required = set(spec.get("required", []))
        rows = [
            parse_column(table, name, meta, required)
            for name, meta in spec.get("properties", {}).items()
        ]
        add_table(doc, ["Cột", "Kiểu", "Ràng buộc", "Ghi chú nghiệp vụ"], rows, [2200, 1700, 2850, 2610])

    doc.add_heading("6. Ghi chú vận hành", level=1)
    for item in [
        "Frontend/API sử dụng Supabase là nguồn dữ liệu chính; Google login đi qua Supabase OAuth.",
        "Backend dùng service role cho nhiều route Express, vì vậy RLS chủ yếu bảo vệ truy cập trực tiếp từ client/PostgREST.",
        "Các bảng user_unlocked_lessons và user_unlocked_chemicals vẫn còn fallback trong User model cũ nhưng không xuất hiện trong schema live public; nguồn tiến độ hiện hành là user_progress.",
        "Khi cần cập nhật schema, đối chiếu ba nguồn: full_clean_schema.sql, OpenAPI live của Supabase và danh sách .from()/rpc() trong code runtime.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("AURUM database description - generated from live Supabase metadata").font.size = Pt(8)

    doc.save(OUT)


if __name__ == "__main__":
    build()
    print(OUT)
